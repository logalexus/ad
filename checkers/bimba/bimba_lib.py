import os
import requests
from checklib import *
import pickle
import random
from dataclasses import dataclass
import json
from bs4 import BeautifulSoup
import re
import docx


PORT = 8070


@dataclass
class Document:
    name: str
    text: str
    data: dict
    uuid: str

    def serialize(self) -> str:
        bts = pickle.dumps(self)
        return bts.hex()

    @classmethod
    def deserialize(cls, raw: str) -> "Document":
        return pickle.loads(bytes.fromhex(raw))

    @classmethod
    def generate_doc(cls):
        name = rnd_string(10)
        text = (
            rnd_string(20)
            + " {{name}} "
            + rnd_string(random.randint(1, 20))
            + " {{age}} "
            + rnd_string(random.randint(1, 20))
        )
        doc = docx.Document()
        doc.add_paragraph(text)
        os.makedirs("/tmp/uploads", exist_ok=True)
        doc.save(f"/tmp/uploads/{name}.docx")
        uuid = ""
        data = {
            "name": rnd_string(random.randint(1, 10)),
            "age": str(random.randint(18, 100)),
        }
        return Document(name, text, data, "")


@dataclass
class User:
    username: str
    password: str
    docs: list

    @classmethod
    def random(cls):
        username = rnd_username()
        password = rnd_password()
        return User(username, password, [])

    def serialize(self) -> str:
        bts = pickle.dumps(self)
        return bts.hex()

    @classmethod
    def deserialize(cls, raw: str) -> "User":
        return pickle.loads(bytes.fromhex(raw))


@dataclass
class Document_card:
    name: str
    uuid: str


def brew_soup(c: BaseChecker, text: str) -> BeautifulSoup:
    try:
        return BeautifulSoup(text, "html.parser")
    except Exception as e:
        c.cquit(Status.MUMBLE, "Failed to parse response", str(e))


class CheckMachine:
    @property
    def url(self):
        return f"http://{self.c.host}:{self.port}"

    def __init__(self, checker: BaseChecker):
        self.c = checker
        self.port = PORT

    def register(self, s: requests.Session, username: str, password: str):
        resp = s.post(
            f"{self.url}/auth/register",
            json={"username": username, "password": password},
        ).json()
        self.c.assert_in(
            resp["message"],
            "User registered successfully!",
            f"Failed to register: {resp}",
        )

    def register_user(self, s: requests.Session, user: User):
        self.register(s, user.username, user.password)

    def login(
        self,
        s: requests.Session,
        user: User,
        status: Status = Status.MUMBLE,
    ):
        s.post(
            f"{self.url}/auth/login",
            json={
                "username": user.username,
                "password": user.password,
            },
        )
        self.c.assert_in("jwt", s.cookies, "Failed to login", status)

    def upload(
        self,
        s: requests.Session,
        document: Document,
        status: Status = Status.MUMBLE,
    ):
        with open(f"/tmp/uploads/{document.name}.docx", "rb") as f:
            file = f.read()
        resp = s.post(
            f"{self.url}/document/upload",
            data={"name": document.name, "object": json.dumps(document.data)},
            files={"file": file},
        ).json()
        self.c.assert_in(resp["message"], "File uploaded successfully!", status)
        return resp["uuid"]

    def get_documents(
        self,
        s: requests.Session,
    ):
        s.get(f"{self.url}/document")

    def search_document(
        self,
        s: requests.Session,
        name: str,
        status: Status = Status.MUMBLE,
    ):
        res = s.get(
            f"{self.url}/document/search?name={name}",
        ).text
        soup = brew_soup(self.c, res)
        try:
            document_item = soup.find("div", class_="document-item")
            name = document_item.find_all("span")
            uuid = re.findall(r"downloadDocument\('(.*?)'", str(document_item))

            return Document_card(name, uuid[0])
        except Exception as e:
            self.c.cquit(
                status, "Failed to parse general information in document search", str(e)
            )

    def download(
        self,
        s: requests.Session,
        uuid: str,
        status: Status = Status.MUMBLE,
    ):
        content = s.get(f"{self.url}/document/download/{uuid}").content
        with open(f"/tmp/{uuid}.docx", "wb") as f:
            f.write(content)

        return content
